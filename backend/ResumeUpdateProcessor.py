from datetime import datetime, timedelta
from typing import Dict, List, Any
import logging
from cosmosdb import CosmosDBManager
import uuid
from azure.search.documents import SearchClient
from azure.search.documents.models import QueryType
from azure.core.credentials import AzureKeyCredential
from openai import AzureOpenAI
from dotenv import load_dotenv
import os
from prompts import generate_work_experience_system_prompt
from azure.core.credentials import AzureKeyCredential
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents import SearchClient
from langchain_openai import AzureChatOpenAI
from azure.communication.email import EmailClient  # Added back

from pydantic import BaseModel

from datetime import datetime, timedelta
from typing import Dict, List, Any
import logging
from cosmosdb import CosmosDBManager
import uuid
from azure.search.documents import SearchClient
from azure.search.documents.models import QueryType
from azure.core.credentials import AzureKeyCredential
from openai import AzureOpenAI
from dotenv import load_dotenv
import os
from azure.storage.blob import BlobServiceClient
from io import BytesIO
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
import contextlib
import tempfile
import re
from prompts import insertion_system_prompt
from langchain_openai import AzureOpenAIEmbeddings
from datetime import timezone
from prompts import generate_work_experience_system_prompt
from azure.core.credentials import AzureKeyCredential
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents import SearchClient
from langchain_openai import AzureChatOpenAI
from azure.communication.email import EmailClient  # Added back
from azure.identity import DefaultAzureCredential

load_dotenv()

# Azure OpenAI
aoai_deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")
aoai_key = os.getenv("AZURE_OPENAI_API_KEY")
aoai_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
aoai_embedding_deployment = os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT_NAME")



primary_llm = AzureChatOpenAI(
    azure_deployment=aoai_deployment,
    api_version="2024-05-01-preview",
    temperature=0.75,
    max_tokens=None,
    timeout=None,
    max_retries=2,
    api_key=aoai_key,
    azure_endpoint=aoai_endpoint
)

API_VERSION = "2024-08-01-preview"

# Azure Blob Storage
connect_str = os.getenv("STORAGE_ACCOUNT_CONNECTION_STRING")
storage_account_name = os.getenv("STORAGE_ACCOUNT_NAME")
container_name = "resumes"


# Initialize Azure OpenAI client
try:
    aoai_client = AzureOpenAI(
        azure_endpoint=aoai_endpoint,
        api_key=aoai_key,
        api_version=API_VERSION
    )
except Exception as e:
    print(f"Failed to initialize Azure OpenAI client: {e}")
    raise

primary_llm_description_json = AzureChatOpenAI(
    azure_deployment=aoai_deployment,
    api_version="2024-08-01-preview",
    temperature=0,
    max_tokens=None,
    timeout=None,
    max_retries=2,
    api_key=aoai_key,
    azure_endpoint=aoai_endpoint,
    model_kwargs={
    "response_format": {"type": "json_schema",
    "json_schema": {
      "name": "project_title_and_description",
      "schema": {
        "type": "object",
        "properties": {
          "title": {
            "type": "string"
          },
          "description": {
            "type": "string"
          }
        },
        "required": [
          "title",
          "description"
        ]
      }
    }
    }
    }
)

primary_llm_insertion_json = AzureChatOpenAI(
    azure_deployment=aoai_deployment,
    api_version="2024-08-01-preview",
    temperature=0,
    max_tokens=None,
    timeout=None,
    max_retries=2,
    api_key=aoai_key,
    azure_endpoint=aoai_endpoint,
    model_kwargs={
    "response_format": {"type": "json_schema",
    "json_schema": {
      "name": "project_insertion_position",
      "schema": {
        "type": "object",
        "properties": {
          "analysis": {
            "type": "string"
          },
          "start_phrase": {
            "type": "string"
          }
        },
        "required": [
          "analysis",
          "start_phrase"
        ]
      }
    }
    }
    }
)

primary_embedding_llm = AzureOpenAIEmbeddings(
    model=aoai_embedding_deployment,
    azure_endpoint=aoai_endpoint,
    api_key=aoai_key,
    openai_api_version="2024-08-01-preview"
)

from typing import List, Dict, Optional, Union
def inference_structured_output_aoai(messages: List[Dict[str, Union[str, List[Dict[str, Union[str, Dict[str, str]]]]]]], deployment: str, schema: BaseModel) -> dict:
    """
    Perform an inference task with structured output using Azure OpenAI.

    Parameters:
    - messages (List[Dict]): A list of message dictionaries. Each dictionary should have a 'role' key
      (either 'system', 'user', or 'assistant') and a 'content' key. The 'content' can be either a string
      for text-only messages or a list of dictionaries for messages that include images.
    - deployment (str): The name of the Azure OpenAI deployment to use.
    - schema (BaseModel): A Pydantic model that defines the structure of the expected output.

    Returns:
    - dict: The parsed response from Azure OpenAI, or None if an error occurs. The response includes
      the generated content structured according to the provided schema.

    This function sends the provided messages to Azure OpenAI and returns the model's response
    parsed according to the given schema. It can handle both text-only inputs and inputs that
    include images, similar to the inference_aoai function.
    """
    try:
        completion = aoai_client.beta.chat.completions.parse(
            model=deployment,
            messages=messages,
            response_format=schema,
        )
        print("Structured output inference completed")
        print(f"Completion content: {completion.choices[0].message.content}")
        print(f"Parsed event: {completion.choices[0].message.parsed}")
        return completion
    except Exception as e:
        print(f"Error in structured output inference: {e}")
        return None


class ProjectExperience(BaseModel):
    """Schema for structured project experience output"""
    project_name: str
    project_experience: str

class ProjectTitleAndDescription(BaseModel):
    """Schema for parsing project title and description"""
    title: str
    description: str

class ResumeUpdateProcessor:
    def __init__(self):
        # Constants
        self.NOTIFICATION_COOLDOWN_HOURS = 24
        self.HOURS_THRESHOLD = 40
        #Azure AI Search
        self.ai_search_credential = AzureKeyCredential(os.environ.get("AZURE_SEARCH_KEY"))
        
        # Initialize Azure AI Search clients
        self.search_client_resumes = SearchClient(
            endpoint=os.environ.get("AZURE_SEARCH_ENDPOINT"),
            index_name=os.environ.get("AZURE_SEARCH_INDEX_RESUMES"),
            credential=self.ai_search_credential
        )
        
        self.search_client_projects = SearchClient(
            endpoint=os.environ.get("AZURE_SEARCH_ENDPOINT"),
            index_name=os.environ.get("AZURE_SEARCH_INDEX_PROJECTS"),
            credential=self.ai_search_credential
        )
        
        self.events_container = CosmosDBManager(
            cosmos_database_id="ResumeAutomation",
            cosmos_container_id="project_key_members"
        )
        self.trackers_container = CosmosDBManager(
            cosmos_database_id="ResumeAutomation",
            cosmos_container_id="resume_trackers"
        )
        self.notification_container = CosmosDBManager(
            cosmos_database_id="ResumeAutomation",
            cosmos_container_id="notifications"
        )
        self.employee_metadata_container = CosmosDBManager(
            cosmos_database_id="ResumeAutomation",
            cosmos_container_id="employee_metadata"
        )
        self.feedback_container = CosmosDBManager(
            cosmos_database_id="ResumeAutomation",
            cosmos_container_id="feedback"
        )
        self.logger = logging.getLogger(__name__)
        self.notification_cooldown = timedelta(hours=self.NOTIFICATION_COOLDOWN_HOURS)
        
        # Initialize Email Client
        connection_string = os.environ.get("COMMUNICATION_SERVICES_CONNECTION_STRING")
        self.email_client = EmailClient.from_connection_string(connection_string)
        self.webapp_url = os.environ.get("WEBAPP_URL")

        self.blob_service_client = None
        try:
            # First try connection string
            connect_str = os.getenv("STORAGE_ACCOUNT_CONNECTION_STRING")
            if connect_str:
                self.logger.info("Initializing blob storage with connection string")
                self.blob_service_client = BlobServiceClient.from_connection_string(connect_str)
            else:
                # Fall back to DefaultAzureCredential
                self.logger.info("No connection string found, using DefaultAzureCredential")
                storage_account_name = os.getenv("STORAGE_ACCOUNT_NAME")
                account_url = f"https://{storage_account_name}.blob.core.windows.net"
                credential = DefaultAzureCredential()
                self.blob_service_client = BlobServiceClient(account_url=account_url, credential=credential)
        except Exception as e:
            self.logger.error(f"Failed to initialize blob storage client: {str(e)}")
            raise

        # Store container names
        self.resume_container_name = os.getenv("STORAGE_ACCOUNT_RESUME_CONTAINER", "resumes")
        self.input_resumes_folder = os.environ.get("STORAGE_ACCOUNT_INPUT_FOLDER", "processed")
        self.updated_resumes_folder = os.environ.get("STORAGE_ACCOUNT_OUTPUT_FOLDER", "updated")
            
    def process_key_member(self, member_entry: Dict) -> Dict:
        """Process an incoming key member entry."""
        try:
            # 1. Store the event
            
            stored_event = self._store_event(member_entry)
            
            # 2. Get or create resume tracker
            tracker = self._get_or_create_tracker(member_entry)
            
            # 3. Check if we need to trigger resume update
            if self._should_trigger_update(tracker):
                # Update tracker status to in_progress first
                tracker['added_to_resume'] = 'in_progress'
                tracker['last_updated'] = datetime.utcnow().isoformat()
                tracker['version'] += 1
                updated_tracker = self.trackers_container.update_item(tracker)
                
                # Then trigger the draft creation
                final_tracker = self._trigger_draft_creation(updated_tracker)
                
                return {
                    'status': 'triggered',
                    'message': f"Resume update triggered. Total hours: {final_tracker['total_hours']}",
                    'event_id': stored_event['id'],
                    'tracker_id': final_tracker['id']
                }
            
            return {
                'status': 'stored',
                'message': f"Event stored, total hours: {tracker['total_hours']}",
                'event_id': stored_event['id'],
                'tracker_id': tracker['id']
            }

        except Exception as e:
            self.logger.error(f"Error processing key member entry: {str(e)}")
            raise

    def _get_resume(self, employee_id: str) -> Dict:
        """Get the resume for the given employee ID."""
        print("Getting resume for employee_id: ", employee_id)
        print(f"Querying {self.search_client_resumes._index_name} for employee_id: ", employee_id)
        results = self.search_client_resumes.search(
            search_text="*",
            filter="employee_id eq '" + employee_id + "'",
            select="id,jobTitle,experienceLevel,content,sourceFileName,employee_id"
        )

  
        for result in results:
            #check if sourceFileName contains employee_id
            #print sourceFileName
            print("SourceFileName: ", result["sourceFileName"])
            return result
        return {}  # added fallback return statement if no results
    
    def _get_project(self, project_number: str) -> Dict:
        search_results = self.search_client_projects.search(
            search_text="*",
            filter="project_number eq '" + project_number + "'",
            select="id, project_number, content, sourcefilename, sourcepage"
        )
        
        sorted_results = sorted(search_results, key=lambda x: x['sourcepage'])
        return sorted_results

    def _generate_project_experience(self, project_data: Dict, role_name: str, resume: Dict) -> Dict:
        """
        Generate a structured project experience using the project data and resume data.
        
        Returns:
            Dict with project_name and project_experience fields
        """
        print("Generating project work experience...")

        project_string = ""
        for project in project_data:
            project_string = project_string + project["content"]
        
        current_date = datetime.utcnow().strftime("%Y-%m-%d")

        generate_work_experience_system_prompt = f"""
            You are an AI assistant that helps with building resumes.
            You will be provided with the candidate's current resume as well as a project description that they recently worked on. 
            Your job will be to use the project description and the format and language of the resume to return a work experience blurb to add into resume.
            The project summary should be in past tense and should be written in a way that is consistent with the rest of the resume. 
            The current date is {current_date}. Estimate the start date of the project based on the current date. It should be <start date> to 'Present'.

            Extract the project name and work experience paragraph. In the work experience, make sure to include the project name, the date range, and the work experience blurb. 
            """
                
        # Prepare messages for LLM
        work_experience_user_message = f"<Current Resume>\n {resume['content']}\n\n <Project Description>\n {project_string}"
        messages = [
            {"role": "system", "content": generate_work_experience_system_prompt},
            {"role": "user", "content": work_experience_user_message}
        ]

        result = inference_structured_output_aoai(messages, aoai_deployment, ProjectExperience)
        if result:
            project_experience = ProjectExperience(**result.choices[0].message.parsed.dict())
            print(f"Project Name: {project_experience.project_name}")
            print(f"Work Experience: {project_experience.project_experience}")

        else:
            print("Failed to process structured output")


        return {
            "project_name": project_experience.project_name,
            "project_experience": project_experience.project_experience
        }

    def _get_employee_email(self, employee_id: str) -> str:
        """Get employee email from metadata container."""
        try:
            results = self.employee_metadata_container.query_items(
                query="SELECT * FROM c WHERE c.employee_id = @employee_id",
                parameters=[{"name": "@employee_id", "value": employee_id}],
                partition_key="metadata"
            )
            
            if results:
                return results[0].get('email')
            else:
                # For testing - return default email if no record exists
                self.logger.warning(f"No email found for employee {employee_id}, using default")
                return "NA"
                
        except Exception as e:
            self.logger.error(f"Error getting employee email: {str(e)}")
            return None

    def _send_notification(self, employee_id: str) -> bool:
        try:
            employee_email = self._get_employee_email(employee_id)
            if not employee_email:
                self.logger.error(f"Could not send notification - no email found for employee {employee_id}")
                return False

            review_link = f"{self.webapp_url}/resume-review"

            message = {
                "senderAddress": "DoNotReply@5fec6054-f6e1-4926-9c37-029ca719c8ae.azurecomm.net",
                "recipients": {
                    "to": [{"address": employee_email}]
                },
                "content": {
                    "subject": "Resume Update Required",
                    "plainText": (
                        "Hello,\n\n"
                        "You have a pending resume update that requires your review. Please take a moment to review and approve these updates to keep your resume current.\n\n"
                        f"Review your updates here: {review_link}\n\n"
                        "This is an automated message. Please do not reply to this email."
                    ),
                    "html": f"""
                    <html>
                        <head>
                            <style>
                                body {{
                                    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Arial, sans-serif;
                                    line-height: 1.6;
                                    margin: 0;
                                    padding: 0;
                                    background-color: #ffffff;
                                }}
                                .container {{
                                    max-width: 600px;
                                    margin: 0 auto;
                                    background: #ffffff;
                                }}
                                .header {{
                                    background-color: #003087;
                                    padding: 20px;
                                    text-align: left;
                                }}
                                .header h1 {{
                                    margin: 0;
                                    color: #ffffff;
                                    font-size: 20px;
                                    font-weight: 500;
                                }}
                                .content {{
                                    padding: 40px 20px;
                                    color: #333333;
                                }}
                                .button {{
                                    display: inline-block;
                                    background-color: #7AC143;
                                    color: #ffffff;
                                    padding: 10px 24px;
                                    text-decoration: none;
                                    border-radius: 4px;
                                    font-weight: 500;
                                    margin-top: 20px;
                                }}
                                .footer {{
                                    padding: 20px;
                                    color: #666666;
                                    font-size: 14px;
                                    border-top: 1px solid #eeeeee;
                                }}
                            </style>
                        </head>
                        <body>
                            <div class="container">
                                <div class="header">
                                    <h1>📄 Resume Update Review Required</h1>
                                </div>
                                <div class="content">
                                    <p>Hello,</p>
                                    <p>You have a pending resume update that requires your review. Please take a moment to review and approve these updates to keep your resume current.</p>
                                    <a href="{review_link}" class="button">Review Updates Now</a>
                                </div>
                                <div class="footer">
                                    This is an automated message. Please do not reply to this email.
                                </div>
                            </div>
                        </body>
                    </html>"""
                }
            }

            poller = self.email_client.begin_send(message)
            result = poller.result()
            self.logger.info(f"Email sent to {employee_email}: {result}")
            return True

        except Exception as e:
            self.logger.error(f"Error sending email notification: {str(e)}")
            return False

    def _update_notification_record(self, employee_id: str) -> Dict:
        """Update notification record with latest notification time or create it if it doesn't exist."""
        try:
            notification_id = f"notification-{employee_id}"
            partition_key = "notifications"
            
            # Check if the notification record exists
            existing_notifications = self.notification_container.query_items(
                query="SELECT * FROM c WHERE c.id = @id",
                parameters=[{"name": "@id", "value": notification_id}],
                partition_key=partition_key
            )
            
            if existing_notifications:
                notification = existing_notifications[0]
                notification['last_notification'] = datetime.utcnow().isoformat()
                return self.notification_container.update_item(notification)
            else:
                # Create new notification record
                notification = {
                    "id": notification_id,
                    "partitionKey": partition_key,
                    "employee_id": employee_id,
                    "last_notification": datetime.utcnow().isoformat()
                }
                return self.notification_container.create_item(notification)

        except Exception as e:
            self.logger.error(f"Error updating notification record: {str(e)}")
            raise
        
    def _is_project_on_resume(self, employee_id: str, project_number: str) -> bool:
        """
        Check if a project is already included in the resume using LLM analysis.
        
        Args:
            employee_id: ID of the employee
            project_number: Project number to check
            
        Returns:
            bool: True if project is found in resume, False otherwise
        """
        try:
            # Get resume
            resume = self._get_resume(employee_id)
            if not resume:
                self.logger.warning(f"No resume found for employee {employee_id}")
                return False
                
            # Get project description
            project_data = self._get_project(project_number)
            if not project_data:
                self.logger.warning(f"No project data found for project {project_number}")
                return False
                
            # Concatenate project content like in _generate_project_experience
            project_string = ""
            for project in project_data:
                project_string = project_string + project["content"]
            
            # Prepare system prompt
            system_prompt = """You are an expert resume analyst. Your task is to determine if a specific project is already mentioned 
            in a resume. Analyze the content carefully and consider that the same project might be described using different words or 
            phrases. Look for matching:
            - Project descriptions
            - Key responsibilities
            - Technologies used
            - Time periods
            - Project outcomes
            
            Respond with only "yes" if you are confident the project is already included in the resume, or "no" if it is not included 
            or if you are unsure. Do not provide any other explanation or commentary."""
            
            # Prepare user message
            user_message = f"""Please analyze if this project is already included in the resume:

            RESUME CONTENT:
            {resume.get('content', '')}
            
            PROJECT DESCRIPTION:
            {project_string}
            
            Is this project already included in the resume? Answer only 'yes' or 'no'."""
            
            # Prepare messages for LLM
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ]
            
            
            # Call LLM
            response = primary_llm.invoke(messages)
            
            # Extract answer and convert to boolean
            answer = response.content.strip().lower()
            print(f"Is project on resume?: {answer}")
            self.logger.info(f"LLM response for project inclusion check: {answer}")
            
            #if answer == "yes", update tracker with added_to_resume = 'yes'
            if answer == "yes":
                #update tracker
                tracker_id = f"{project_number}-{employee_id}"
                tracker = self.trackers_container.query_items(
                    query="SELECT * FROM c WHERE c.id = @id",
                    parameters=[{"name": "@id", "value": tracker_id}],
                    partition_key="resumeupdatestatus"
                )
                if tracker:
                    tracker[0]['added_to_resume'] = 'yes'
                    tracker[0]['last_updated'] = datetime.utcnow().isoformat()
                    tracker[0]['version'] += 1
                    self.trackers_container.update_item(tracker[0])
                    print("Tracker updated with added_to_resume='yes'")
            
            return answer == "yes"
            
        except Exception as e:
            self.logger.error(f"Error checking if project is on resume: {str(e)}")
            return False

    def _trigger_draft_creation(self, tracker: Dict) -> Dict:
        """Process resume update for the given tracker."""
        try:
            # Get current resume
            resume = self._get_resume(tracker['employee_id'])
            
            # Get project details
            project = self._get_project(tracker['project_number'])
            
            # Get current role from tracker's role history
            current_role = tracker['role_history'][-1]['role_name']
            
            # Generate project experience with structured output
            generated_content = self._generate_project_experience(
                project_data=project,
                role_name=current_role,
                resume=resume
            )
            
            # Update tracker with the generated content
            tracker['project_name'] = generated_content['project_name']
            tracker['description'] = generated_content['project_experience']
            tracker['last_updated'] = datetime.utcnow().isoformat()
            tracker['version'] += 1
            
            # Save the updated tracker
            updated_tracker = self.trackers_container.update_item(tracker)

            # Check notification cooldown before sending notification
            if self._should_send_notification(tracker['employee_id']):
                if self._send_notification(tracker['employee_id']):
                    print("Notification sent successfully")
                    self._update_notification_record(tracker['employee_id'])
                else:
                    print("Failed to send notification")
            else:
                print(f"Skipping notification for employee {tracker['employee_id']} due to cooldown period")
            
            return updated_tracker
            
        except Exception as e:
            self.logger.error(f"Error in trigger_draft_creation: {str(e)}")
            raise
        
    def _should_trigger_update(self, tracker: Dict) -> bool:
        """
        Determine if we should trigger a resume update.
        Returns True if hours >= 40 and status is 'no'
        """
        #check if total_hours >=40 and added_to_resume == 'no'
        print("Checking if we should trigger update...")
        if tracker['total_hours'] >= 40 and tracker['added_to_resume'] == 'no':
            # Extract employee_id and project_number from the tracker
            employee_id = tracker['employee_id']
            project_number = tracker['project_number']
            

            print("Hours > 40 and added_to_resume='no'. Checking if project is on resume...")
            # Run _is_project_on_resume() to check if the project is on the resume
            if self._is_project_on_resume(employee_id, project_number):
                #Set added_to_resume to 'yes' and return False


                return False
        
            return True

        return False  

    def _store_event(self, member_entry: Dict) -> Dict:
        """Store the key member entry in the events container."""
        # Extract employee ID from display name (e.g., "Diaz, John - 18117" -> "18117")
        employee_id = member_entry['employee_display_name'].split(' - ')[1]
        
        event_doc = {
            "id": str(uuid.uuid4()),
            "partitionKey": member_entry['project_number'],  # Use project number as partition key
            "type": "project_key_member",
            "subject_area": member_entry['subject_area'],
            "project_number": member_entry['project_number'],
            "project_role_name": member_entry['project_role_name'],
            "employee_display_name": member_entry['employee_display_name'],
            "employee_id": employee_id,
            "job_hours": member_entry['job_hours'],
            "employee_job_family_function_code": member_entry['employee_job_family_function_code'],
            "timestamp": datetime.utcnow().isoformat()
        }
        print(f'Storing event for Employee_id={employee_id} and project_number={member_entry["project_number"]}')
        return self.events_container.create_item(event_doc)

    def _get_or_create_tracker(self, member_entry: Dict) -> Dict:
        """Get or create tracker in the trackers container."""
        # Extract employee ID from display name
        employee_id = member_entry['employee_display_name'].split(' - ')[1]
        
        # Create compound ID from project and employee
        tracker_id = f"{member_entry['project_number']}-{employee_id}"
        
        try:
            # Try to get existing tracker directly by ID
            tracker = self.trackers_container.query_items(
                query="SELECT * FROM c WHERE c.id = @id",
                parameters=[{"name": "@id", "value": tracker_id}],
                partition_key="resumeupdatestatus"
            )
            
            if tracker:
                existing_tracker = tracker[0]
                # Update with latest hours - no need to sum up events anymore
                existing_tracker['total_hours'] = member_entry['job_hours']
                existing_tracker['last_updated'] = datetime.utcnow().isoformat()
                existing_tracker['version'] += 1
                
                # Update role history if needed
                self._update_role_history(existing_tracker, member_entry)
                
                return self.trackers_container.update_item(existing_tracker)
            
            # Create new tracker if not found
            new_tracker = {
                "id": tracker_id,
                "partitionKey": "resumeupdatestatus",
                "type": "resume_update_tracker",
                "employee_id": employee_id,
                "employee_display_name": member_entry['employee_display_name'],
                "project_number": member_entry['project_number'],
                "subject_area": member_entry['subject_area'],
                "total_hours": member_entry['job_hours'],
                "added_to_resume": "no",
                "project_name": "",  # Added new field
                "description": "",
                "created_timestamp": datetime.utcnow().isoformat(),
                "last_updated": datetime.utcnow().isoformat(),
                "version": 1,
                "role_history": [
                    {
                        "role_name": member_entry['project_role_name'],
                        "job_family_code": member_entry['employee_job_family_function_code'],
                        "start_date": datetime.utcnow().isoformat(),
                        "end_date": None
                    }
                ],
                "review_date": (datetime.utcnow() + timedelta(days=30)).isoformat()
            }
            return self.trackers_container.create_item(new_tracker)
            
        except Exception as e:
            self.logger.error(f"Error in get_or_create_tracker: {str(e)}")
            raise

    def _update_role_history(self, tracker: Dict, member_entry: Dict):
        """Update role history if role has changed."""
        current_role = tracker['role_history'][-1]
        if (current_role['role_name'] != member_entry['project_role_name'] or 
            current_role['job_family_code'] != member_entry['employee_job_family_function_code']):
            # Close current role
            current_role['end_date'] = datetime.utcnow().isoformat()
            # Add new role
            tracker['role_history'].append({
                "role_name": member_entry['project_role_name'],
                "job_family_code": member_entry['employee_job_family_function_code'],
                "start_date": datetime.utcnow().isoformat(),
                "end_date": None
            })

    def get_pending_updates(self, employee_id: str = None) -> List[Dict]:
        """
        Get all pending resume updates for a specific employee or all employees.
        Returns updates where added_to_resume = 'in_progress'.
        
        Args:
            employee_id: Optional employee ID. If None, returns updates for all employees.
            
        Returns:
            List[Dict]: List of pending resume update trackers
        """
        base_query = """
            SELECT *
            FROM c
            WHERE c.partitionKey = 'resumeupdatestatus'
            AND c.added_to_resume = 'in_progress'
        """
        
        if employee_id:
            query = base_query + " AND c.employee_id = @employee_id"
            parameters = [{"name": "@employee_id", "value": employee_id}]
        else:
            query = base_query
            parameters = []
        
        return self.trackers_container.query_items(
            query=query,
            parameters=parameters,
            partition_key="resumeupdatestatus"
        )
    
    def _parse_project_title_and_description(self, full_description: str) -> Dict[str, str]:
        """
        Use LLM to separate the title and description from a project description.
        
        Args:
            full_description: The complete project description
            
        Returns:
            Dict with 'title' and 'description' keys
        """
        try:
            system_prompt = """You are an expert at parsing resume entries. Given a project description 
            that starts with a title line followed by details, separate out the title line from the description.
            The title line typically includes the role, project name, location, and date range.

            Make sure:
            1. The title includes ALL parts (role, project name, location, dates)
            2. The description does not repeat the title information
            3. Both sections maintain proper grammar and formatting"""

            user_message = f"Please parse this project entry:\n\n{full_description}"
            
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ]
            
            result = inference_structured_output_aoai(messages, aoai_deployment, ProjectTitleAndDescription)
            if result:
                parsed_content = ProjectTitleAndDescription(**result.choices[0].message.parsed.dict())
                print(f"Parsed title: {parsed_content.title}")
                print(f"Parsed description: {parsed_content.description}")
                return {
                    "title": parsed_content.title,
                    "description": parsed_content.description
                }
            else:
                self.logger.error("Failed to parse project content")
                return None

        except Exception as e:
            self.logger.error(f"Error parsing project title and description: {str(e)}")
            return None

    def save_updates(self, employee_id: str, projects: List[Dict]) -> bool:
        """
        Save multiple resume updates by updating their status and modifying the resume document.
        Updates trackers only after successful document save to maintain consistency.
        
        Args:
            employee_id: ID of the employee
            projects: List of project dictionaries containing project_number and description
            
        Returns:
            bool: True if all updates were saved successfully, False otherwise
        """
        try:
            self.logger.info(f"Saving updates for employee {employee_id}, projects: {projects}")
            
            # Get the employee's current resume
            resume = self._get_resume(employee_id)
            if not resume:
                self.logger.error(f"No resume found for employee {employee_id}")
                return False
            
            resume_name = resume.get('sourceFileName')
            print(f"Resume name: {resume_name}")
            if not resume_name:
                self.logger.error(f"No resume filename found for employee {employee_id}")
                return False

            # Download and prepare resume document once
            doc = self._get_resume_document(resume_name)
            if not doc:
                self.logger.error(f"Failed to download resume document: {resume_name}")
                return False

            # Find insertion point once
            insert_phrase = self._find_insert_position(doc)
            if not insert_phrase:
                self.logger.error("Failed to find insertion point in resume")
                return False

            # Collect trackers to update
            trackers_to_update = []

            # Process each project update
            for project in projects:
                try:
                    project_number = project['project_number']
                    description = project['description']
                    
                    # Parse the title and description
                    parsed_content = self._parse_project_title_and_description(description)
                    if not parsed_content:
                        self.logger.error(f"Failed to parse content for project {project_number}")
                        return False

                    # Get tracker
                    tracker_id = f"{project_number}-{employee_id}"
                    results = self.trackers_container.query_items(
                        query="SELECT * FROM c WHERE c.id = @id",
                        parameters=[{"name": "@id", "value": tracker_id}],
                        partition_key="resumeupdatestatus"
                    )
                    
                    results_list = list(results)
                    if not results_list:
                        self.logger.error(f"No tracker found with ID: {tracker_id}")
                        continue
                            
                    tracker = results_list[0]
                    
                    # Prepare tracker update (but don't save yet)
                    tracker['project_name'] = parsed_content['title']  # Store the parsed title
                    tracker['description'] = parsed_content['description']  # Store the parsed description
                    tracker['added_to_resume'] = 'yes'
                    tracker['last_updated'] = datetime.utcnow().isoformat()
                    tracker['version'] += 1
                    
                    trackers_to_update.append(tracker)
                    
                    # Add project to document using parsed content
                    doc = self._save_new_project(doc, json.dumps(parsed_content), insert_phrase)
                    
                except Exception as e:
                    self.logger.error(f"Error processing project {project_number}: {str(e)}")
                    return False

            # Save updated resume once
            enhanced_docx_name = self._save_resume_document(doc, resume_name, resume)
            if not enhanced_docx_name:
                self.logger.error("Failed to save updated resume")
                return False

            # Only update trackers after successful document save
            for tracker in trackers_to_update:
                try:
                    self.trackers_container.update_item(tracker)
                    self.logger.info(f"Successfully updated tracker: {tracker['id']}")
                except Exception as e:
                    self.logger.error(f"Error updating tracker {tracker['id']}: {str(e)}")
                    # Log error but continue with other trackers
                    continue

            self.logger.info(f"Successfully updated resume: {enhanced_docx_name}")
            return True
                    
        except Exception as e:
            self.logger.error(f"Error saving updates: {str(e)}")
            return False
            


    def _get_resume_document(self, resume_name: str) -> Document:
        """Download resume document from blob storage."""
        try:
            print("Downloading resume document...")
            
            blob_name = f"{self.input_resumes_folder}/{resume_name}"
            print(f"Downloading resume document: {blob_name}")
            
            container_client = self.blob_service_client.get_container_client(self.resume_container_name)
            blob_client = container_client.get_blob_client(blob_name)

            if not blob_client.exists():
                self.logger.error(f"Blob {blob_name} not found in {self.resume_container_name}.")
                return None

            # Download and return document
            blob_data = blob_client.download_blob().readall()
            return Document(BytesIO(blob_data))

        except Exception as e:
            self.logger.error(f"Error downloading resume document: {str(e)}")
            return None

    def _save_resume_document(self, doc: Document, resume_name: str, resume: dict) -> str:
        """Save updated resume document to blob storage and update index."""
        try:
            resume_name_without_ext = os.path.splitext(resume_name)[0]
            enhanced_docx_name = f"{resume_name_without_ext}.docx"

            # Update search index
            print("Converting document to string...")
            document_content = self._read_docx_to_string(doc)
            print("Updating resume index...")
            self._update_resume_index(resume, document_content)
            print("Resume index updated.")

            # Save document to blob storage
            container_client = self.blob_service_client.get_container_client(self.resume_container_name)

            # Save the enhanced DOCX to blob with updated folder path
            with self._temporary_file(suffix='.docx') as temp_docx_path:
                doc.save(temp_docx_path)
                enhanced_docx_blob_name = f"{self.updated_resumes_folder}/{enhanced_docx_name}"
                
                enhanced_docx_client = container_client.get_blob_client(enhanced_docx_blob_name)
                with open(temp_docx_path, "rb") as docx_file:          
                    enhanced_docx_client.upload_blob(docx_file, overwrite=True)

                print(f"Enhanced resume (DOCX) uploaded as {enhanced_docx_blob_name}")
            return enhanced_docx_blob_name

        except Exception as e:
            self.logger.error(f"Error saving resume document: {str(e)}")
            return None



    #NEW
    def _read_docx_to_string(self, doc: Document):
        # Extract text from each paragraph and concatenate into a single string
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        return '\n'.join(full_text)

    #NEW
    @contextlib.contextmanager
    def _temporary_file(self, suffix=None):
        """Context manager for creating and cleaning up a temporary file."""
        fd, path = tempfile.mkstemp(suffix=suffix)
        try:
            os.close(fd)
            yield path
        finally:
            try:
                os.unlink(path)
            except OSError:
                pass

    #NEW            
    def _find_insert_position(self, doc):
        # Extract text from the document
        print("Finding insert position...")
        full_text = "\n".join([para.text for para in doc.paragraphs])
        # Prompt for the LLM

        messages = [
            {"role": "system", "content": insertion_system_prompt},
            {"role": "user", "content": full_text}
        ]
        
        result = primary_llm_insertion_json.invoke(messages)
        result_json = json.loads(result.content)
        
        print("Analysis:", result_json['analysis'])
        print("Start Phrase:", result_json['start_phrase'])

        return result_json['start_phrase']
    
    def _save_new_project(self, doc: Document, new_project: str, insert_phrase: str):

        print(f"Inserting new project: {new_project}")
        json_project = json.loads(new_project)
        for para in doc.paragraphs:
            if insert_phrase in para.text:
                # Insert the new project before the paragraph containing the insert phrase
                new_para = para.insert_paragraph_before()
                new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = new_para.add_run(json_project['title'])
                run.bold = True
                run.font.size = Pt(10)
                new_para.add_run('. ' + json_project['description'])
                new_para.style = 'Normal'
                # Add a blank line after the new project
                para.insert_paragraph_before()
                return doc
        
        return doc
    
    def _update_resume_index(self, resume: dict, resume_content: str):
        """Update the search index with the new resume content."""
        try:
            search_client = SearchClient(
                endpoint=os.environ.get("AZURE_SEARCH_ENDPOINT"),
                index_name=os.environ.get("AZURE_SEARCH_INDEX_RESUMES"),
                credential=self.ai_search_credential  # Using instance variable instead of global
            )

            current_time_iso = datetime.now(timezone.utc).isoformat()
            resume["content"] = resume_content
            resume["searchVector"] = primary_embedding_llm.embed_query(resume_content)
            resume["date"] = current_time_iso
            result = search_client.upload_documents(documents=[resume])
            return True
            
        except Exception as e:
            self.logger.error(f"Error updating resume index: {str(e)}")
            return False
    
    

    def discard_update(self, employee_id: str, project_number: str) -> bool:
        """
        Discard a pending resume update by setting its status to 'discarded'.
        
        Args:
            employee_id: ID of the employee
            project_number: Project number to discard
                
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Create the tracker ID using the combination
            tracker_id = f"{project_number}-{employee_id}"
            self.logger.info(f"Discarding update for tracker ID: {tracker_id}")
            
            # Query the tracker directly by ID
            results = self.trackers_container.query_items(
                query="SELECT * FROM c WHERE c.id = @id",
                parameters=[{"name": "@id", "value": tracker_id}],
                partition_key="resumeupdatestatus"
            )
            
            results_list = list(results)
            if not results_list:
                self.logger.error(f"No tracker found with ID: {tracker_id}")
                return False
                
            tracker = results_list[0]
                    
            # Update the status to 'discarded'
            tracker['added_to_resume'] = 'discarded'
            tracker['last_updated'] = datetime.utcnow().isoformat()
            tracker['version'] += 1
            
            # Save the updated tracker
            self.trackers_container.update_item(tracker)
            
            self.logger.info(f"Successfully discarded update for tracker: {tracker_id}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error discarding update: {str(e)}")
            return False
    
    def recurring_notification(self) -> Dict[str, Any]:
        """
        Process recurring notifications for pending resume updates.
        Returns summary of notifications sent.
        """
        try:
            self.logger.info("Starting recurring notification process")
            summary = {
                "total_employees_processed": 0,  # Added this field
                "notifications_sent": 0,
                "skipped_cooldown": 0,
                "errors": 0
            }

            # Get all employees with pending updates
            pending_updates = self.get_pending_updates()
            
            # Get unique employee IDs
            employees_to_notify = {update['employee_id'] for update in pending_updates}
            summary["total_employees_processed"] = len(employees_to_notify)
            
            # Process each employee
            for employee_id in employees_to_notify:
                try:
                    if self._should_send_notification(employee_id):
                        if self._send_notification(employee_id):
                            self._update_notification_record(employee_id)
                            summary["notifications_sent"] += 1
                    else:
                        summary["skipped_cooldown"] += 1
                except Exception as e:
                    self.logger.error(f"Error processing notifications for employee {employee_id}: {str(e)}")
                    summary["errors"] += 1

            self.logger.info(f"Recurring notification process completed: {summary}")
            return summary

        except Exception as e:
            self.logger.error(f"Error in recurring notification process: {str(e)}")
            raise


    def _should_send_notification(self, employee_id: str) -> bool:
        """
        Check if notification should be sent based on cooldown period.
        Returns True if last notification was more than cooldown period ago.
        """
        try:
            query = """
                SELECT * FROM c 
                WHERE c.partitionKey = 'notifications'
                AND c.id = @notification_id
            """
            parameters = [
                {"name": "@notification_id", "value": f"notification-{employee_id}"}
            ]
            results = list(self.notification_container.query_items(
                query=query,
                parameters=parameters,
                partition_key="notifications"
            ))

            if not results:
                return True

            notification_record = results[0]
            last_notification = datetime.fromisoformat(notification_record['last_notification'])
            
            return datetime.utcnow() - last_notification > self.notification_cooldown

        except Exception as e:
            self.logger.error(f"Error checking notification eligibility: {str(e)}")
            return False

        except Exception as e:
            self.logger.error(f"Error checking notification eligibility: {str(e)}")
            return False

    def store_feedback(self, feedback_data: dict) -> dict:
        """
        Store feedback in the feedback container.
        
        Args:
            feedback_data (dict): Dictionary containing feedback type and content
            
        Returns:
            dict: The stored feedback document
        """
        try:
            employee_id = feedback_data.get('employee_id')
            
            feedback_doc = {
                "id": f"feedback-{datetime.utcnow().strftime('%Y%m%d-%H%M%S')}-{employee_id}",
                "partitionKey": "AutoResumeUpdates",  # Fixed partition key as requested
                "type": "feedback",
                "feedback_type": feedback_data.get('type'),
                "content": feedback_data.get('content'),
                "employee_id": employee_id,
                "created_timestamp": datetime.utcnow().isoformat(),
                "status": "new"
            }
            
            stored_feedback = self.feedback_container.create_item(feedback_doc)
            self.logger.info(f"Stored feedback document with id: {stored_feedback['id']}")
            
            return stored_feedback
            
        except Exception as e:
            self.logger.error(f"Error storing feedback: {str(e)}")
            raise
    

    def reset_resume(self, employee_id: str) -> bool:
        """
        Reset the search index to match the original resume in the /processed folder.
        
        Args:
            employee_id: ID of the employee
            
        Returns:
            bool: True if reset was successful, False otherwise
        """
        try:
            # Get the resume metadata from search index
            resume = self._get_resume(employee_id)
            if not resume:
                self.logger.error(f"No resume found for employee {employee_id}")
                return False
            
            resume_name = resume.get('sourceFileName')
            if not resume_name:
                self.logger.error(f"No resume filename found for employee {employee_id}")
                return False

            # Download the original resume from /processed folder
            doc = self._get_resume_document(resume_name)
            if not doc:
                self.logger.error(f"Failed to download original resume: {resume_name}")
                return False

            # Convert document to string
            document_content = self._read_docx_to_string(doc)
            
            # Update the search index with original content
            if self._update_resume_index(resume, document_content):
                self.logger.info(f"Successfully reset search index for employee {employee_id}")
                return True
            else:
                self.logger.error(f"Failed to update search index for employee {employee_id}")
                return False

        except Exception as e:
            self.logger.error(f"Error resetting resume: {str(e)}")
            return False